import streamlit as st
from llama_index.core import VectorStoreIndex, Document
from llama_index.llms.openai import OpenAI
from llama_index.core import Settings
import os
import pdfplumber
from docx import Document as DocxDocument

import json

st.header("Chat with the Streamlit docs 💬 📚")

# Sidebar for OpenAI API Key
openai_api_key = st.sidebar.text_input("Enter your OpenAI API Key:", type="password")

# Initialize session state for messages
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": "Ask me a question about Streamlit's open-source Python library!"}
    ]

# Function to read PDF files
def read_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    return text

# Function to read DOCX files
def read_docx(file):
    doc = DocxDocument(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

@st.cache_resource(show_spinner=False)
def load_data(uploaded_files):
    with st.spinner("Loading and indexing the documents – hang tight! This should take 1-2 minutes."):
        docs = []
        for uploaded_file in uploaded_files:
            if uploaded_file.type == "application/pdf":
                text = read_pdf(uploaded_file)
                docs.append(Document(text=text))
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = read_docx(uploaded_file)
                docs.append(Document(text=text))

        Settings.llm = OpenAI(model="gpt-3.5-turbo", temperature=0.5, 
                              system_prompt="You are an expert on the Streamlit Python library and your job is to answer technical questions. Assume that all questions are related to the Streamlit Python library. Keep your answers technical and based on facts – do not hallucinate features.")
        
        index = VectorStoreIndex.from_documents(docs, settings=Settings.llm)
        return index

# Function to save the conversation
def save_conversation():
    with open("conversations.json", "a") as f:
        json.dump(st.session_state.messages, f)
        f.write("\n")

# Function to load previous conversations
def load_conversations():
    if os.path.exists("conversations.json"):
        with open("conversations.json", "r") as f:
            conversations = [json.loads(line) for line in f]
        return conversations
    return []

# Function to delete selected conversations
def delete_selected_conversations(selected_indices):
    conversations = load_conversations()
    remaining_conversations = [conv for i, conv in enumerate(conversations) if i not in selected_indices]
    with open("conversations.json", "w") as f:
        for conv in remaining_conversations:
            json.dump(conv, f)
            f.write("\n")

# File uploader for multiple PDF and DOCX files
uploaded_files = st.file_uploader("Upload PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files and st.session_state.openai_api_key:
    index = load_data(uploaded_files)
    chat_engine = index.as_chat_engine(chat_mode="condense_question", verbose=True)

    # User input for questions
    if prompt := st.chat_input("Your question"):
        st.session_state.messages.append({"role": "user", "content": prompt})

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    if len(st.session_state.messages) > 0 and st.session_state.messages[-1]["role"] != "assistant":
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = chat_engine.chat(prompt)
                st.write(response.response)
                message = {"role": "assistant", "content": response.response}
                st.session_state.messages.append(message)

    if st.button("Save Conversation"):
        if st.session_state.messages:
            st.session_state.confirm_save = True

    if st.session_state.get('confirm_save', False):
        st.warning("Do you want to save the conversation?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                save_conversation()
                st.success("Conversation saved!")
                st.session_state.confirm_save = False
        with col2:
            if st.button("No"):
                st.session_state.confirm_save = False

    if st.button("End Conversation"):
        st.session_state.messages = []
        st.success("Conversation ended. You can start a new one!")

else:
    st.sidebar.warning("Please enter your OpenAI API key and upload PDF or DOCX files to proceed.")

# Sidebar to toggle visibility of previous conversations
if 'show_conversations' not in st.session_state:
    st.session_state.show_conversations = False

if st.sidebar.button("Toggle Previous Conversations"):
    st.session_state.show_conversations = not st.session_state.show_conversations

# Show previous conversations if the toggle is enabled
if st.session_state.show_conversations:
    st.sidebar.subheader("Previous Conversations")
    conversations = load_conversations()

    if conversations:
        selected_indices = []
        for i, conv in enumerate(conversations):
            st.sidebar.write(f"Conversation {i + 1}:")
            for message in conv:
                st.sidebar.write(f"{message['role']}: {message['content']}")
            # Checkbox for selecting conversation to delete
            if st.sidebar.checkbox(f"Select Conversation {i + 1} for Deletion", key=f"delete_checkbox_{i}"):
                selected_indices.append(i)

        if st.sidebar.button("Delete Selected Conversations"):
            if selected_indices:
                delete_selected_conversations(selected_indices)
                st.success("Selected conversations deleted. Please Refresh to See the Effect!")
                st.session_state.messages = []  # Optional: reset messages for a fresh start

    else:
        st.sidebar.write("No previous conversations found.")
else:
    st.sidebar.write("Previous conversations are hidden. Click 'Toggle Previous Conversations' to show.")
